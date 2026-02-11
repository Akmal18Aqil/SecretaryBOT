import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_DIR = 'templates'

def create_header(document):
    header = document.sections[0].header
    p = header.paragraphs[0]
    p.text = "PESANTREN MULTIMEDIA CREATIVE SWARM"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.bold = True
    p.style.font.size = Pt(16)
    
    p2 = header.add_paragraph()
    p2.text = "Jl. Coding Tanpa Batas No. 00, Server Localhost, Cloud"
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.style.font.size = Pt(10)
    
    p3 = header.add_paragraph()
    p3.text = "="*60
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_undangan():
    filename = 'tpl_undangan_internal.docx'
    document = Document()
    create_header(document)
    
    document.add_heading('SURAT UNDANGAN', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Nomor: {'{{ nomor_surat }}'}")
    document.add_paragraph(f"Lampiran: -")
    
    document.add_paragraph("Kepada Yth.")
    p_to = document.add_paragraph("{{ penerima }}")
    p_to.style.font.bold = True
    document.add_paragraph("di Tempat")
    document.add_paragraph("") 
    
    document.add_paragraph("Assalamu'alaikum Wr. Wb.")
    document.add_paragraph("Dengan hormat,")
    document.add_paragraph("Mengharap kehadiran Antum pada acara {{ acara }} yang akan dilaksanakan pada:")
    
    p_date = document.add_paragraph()
    p_date.add_run("Hari/Tanggal\t: ").bold = True
    p_date.add_run("{{ hari_tanggal }}")
    
    p_time = document.add_paragraph()
    p_time.add_run("Waktu\t\t: ").bold = True
    p_time.add_run("{{ waktu }}")
    
    p_place = document.add_paragraph()
    p_place.add_run("Tempat\t\t: ").bold = True
    p_place.add_run("{{ tempat }}")
    
    document.add_paragraph("")
    document.add_paragraph("Demikian undangan ini kami sampaikan. Atas perhatian dan kehadirannya kami ucapkan terima kasih.")
    document.add_paragraph("Wassalamu'alaikum Wr. Wb.")
    
    document.add_paragraph("")
    sig = document.add_paragraph("Sekretaris Divisi,")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph("")
    document.add_paragraph("")
    sig_name = document.add_paragraph("( ____________________ )")
    sig_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = os.path.join(TEMPLATE_DIR, filename)
    document.save(path)
    print(f"[SUCCESS] Template Undangan dibuat: {path}")

def create_peminjaman():
    filename = 'tpl_peminjaman_barang.docx'
    document = Document()
    create_header(document)
    
    document.add_heading('SURAT PEMINJAMAN BARANG', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Nomor: {'{{ nomor_surat }}'}")
    
    document.add_paragraph("Kepada Yth.")
    p_to = document.add_paragraph("Bagian Sarana Prasarana")
    p_to.style.font.bold = True
    document.add_paragraph("di Tempat")
    document.add_paragraph("") 
    
    document.add_paragraph("Assalamu'alaikum Wr. Wb.")
    document.add_paragraph("Saya yang bertanda tangan di bawah ini:")
    
    p_nama = document.add_paragraph()
    p_nama.add_run("Nama Pemohon\t: ").bold = True
    p_nama.add_run("{{ pemohon }}") # AI must output 'pemohon'
    
    document.add_paragraph("Bermaksud meminjam barang inventaris untuk keperluan {{ keperluan }}, rincian sebagai berikut:")
    
    # Simple list for now
    p_barang = document.add_paragraph()
    p_barang.add_run("Barang\t\t: ").bold = True
    p_barang.add_run("{{ nama_barang }}")
    
    p_waktu = document.add_paragraph()
    p_waktu.add_run("Waktu Pakai\t: ").bold = True
    p_waktu.add_run("{{ waktu_pinjam }}")

    document.add_paragraph("")
    document.add_paragraph("Saya bertanggung jawab penuh atas keamanan barang tersebut.")
    document.add_paragraph("Wassalamu'alaikum Wr. Wb.")
    
    document.add_paragraph("")
    sig = document.add_paragraph("Pemohon,")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph("")
    document.add_paragraph("")
    sig_name = document.add_paragraph("( {{ pemohon }} )")
    sig_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = os.path.join(TEMPLATE_DIR, filename)
    document.save(path)
    print(f"[SUCCESS] Template Peminjaman dibuat: {path}")

def create_peminjaman_ruangan():
    filename = 'tpl_peminjaman_ruangan.docx'
    document = Document()
    create_header(document)
    
    document.add_heading('SURAT PEMINJAMAN RUANGAN', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Nomor: {'{{ nomor_surat }}'}")
    
    document.add_paragraph("Kepada Yth.")
    p_to = document.add_paragraph("Kepala Bagian Umum")
    p_to.style.font.bold = True
    document.add_paragraph("di Tempat")
    document.add_paragraph("") 
    
    document.add_paragraph("Assalamu'alaikum Wr. Wb.")
    document.add_paragraph("Saya yang bertanda tangan di bawah ini:")
    
    p_nama = document.add_paragraph()
    p_nama.add_run("Nama Pemohon\t: ").bold = True
    p_nama.add_run("{{ pemohon }}") 
    
    document.add_paragraph("Bermaksud meminjam ruangan untuk kegiatan {{ acara }}, dengan rincian sebagai berikut:")
    
    p_ruang = document.add_paragraph()
    p_ruang.add_run("Ruangan\t\t: ").bold = True
    p_ruang.add_run("{{ nama_ruangan }}")

    p_tanggal = document.add_paragraph()
    p_tanggal.add_run("Hari/Tanggal\t: ").bold = True
    p_tanggal.add_run("{{ hari_tanggal }}")
    
    p_waktu = document.add_paragraph()
    p_waktu.add_run("Waktu\t\t: ").bold = True
    p_waktu.add_run("{{ waktu_mulai }} s/d {{ waktu_selesai }}")

    document.add_paragraph("")
    document.add_paragraph("Kami bertanggung jawab penuh atas kebersihan dan kerapihan ruangan setelah digunakan.")
    document.add_paragraph("Wassalamu'alaikum Wr. Wb.")
    
    document.add_paragraph("")
    sig = document.add_paragraph("Pemohon,")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph("")
    document.add_paragraph("")
    sig_name = document.add_paragraph("( {{ pemohon }} )")
    sig_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    path = os.path.join(TEMPLATE_DIR, filename)
    document.save(path)
    print(f"[SUCCESS] Template Peminjaman Ruangan dibuat: {path}")

def main():
    if not os.path.exists(TEMPLATE_DIR):
        os.makedirs(TEMPLATE_DIR)
    
    create_undangan()
    create_peminjaman()
    create_peminjaman_ruangan()

if __name__ == "__main__":
    main()
