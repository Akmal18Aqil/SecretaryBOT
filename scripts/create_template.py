from docx import Document

def create_notulensi_template():
    doc = Document()
    doc.add_heading('NOTULENSI RAPAT', 0)

    doc.add_paragraph('Hari/Tanggal\t: {{ hari_tanggal }}')
    doc.add_paragraph('Waktu\t\t: {{ waktu }}')
    doc.add_paragraph('Tempat\t\t: {{ tempat }}')
    doc.add_paragraph('Agenda\t\t: {{ agenda }}')

    doc.add_heading('Pembahasan:', level=1)
    doc.add_paragraph('{{ pembahasan }}')
    # Note: Complex lists usually need jinja loops inside docx, 
    # but for simple listening let's just dump the text block first or assume the AI formats it as a bulleted string.

    doc.add_heading('Kesimpulan:', level=1)
    doc.add_paragraph('{{ kesimpulan }}')

    doc.add_heading('Tugas / Action Items:', level=1)
    doc.add_paragraph('{{ tugas }}')

    doc.save('templates/tpl_notulensi.docx')
    print("Template templates/tpl_notulensi.docx created.")

if __name__ == "__main__":
    create_notulensi_template()
